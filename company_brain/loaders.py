from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from company_brain.models import ExtractedSection

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".docx"}

logger = logging.getLogger(__name__)


def iter_supported_files(root: Path) -> list[Path]:
    if not root.exists():
        raise FileNotFoundError(f"Input path does not exist: {root}")
    if root.is_file():
        return [root] if root.suffix.lower() in SUPPORTED_EXTENSIONS else []
    return sorted(
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def extract_sections(path: Path) -> list[ExtractedSection]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".xlsx":
        return _read_xlsx(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def extract_text(path: Path) -> str:
    return "\n\n".join(section.content for section in extract_sections(path))


def _read_pdf(path: Path) -> list[ExtractedSection]:
    reader = PdfReader(str(path))
    sections: list[ExtractedSection] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
            if text.strip():
                sections.append(
                    ExtractedSection(
                        content=text.strip(),
                        page_number=index,
                        metadata={"source_type": "pdf_page"},
                    )
                )
        except Exception as exc:  # pragma: no cover - defensive against malformed PDFs
            logger.warning("Failed to extract page %s from %s: %s", index, path, exc)
    return sections


def _read_xlsx(path: Path) -> list[ExtractedSection]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[ExtractedSection] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None]
            if values:
                rows.append(" | ".join(values))
        if rows:
            sections.append(
                ExtractedSection(
                    content="\n".join(rows),
                    sheet_name=sheet.title,
                    heading=sheet.title,
                    metadata={"source_type": "xlsx_sheet", "row_count": len(rows)},
                )
            )
    workbook.close()
    return sections


def _read_docx(path: Path) -> list[ExtractedSection]:
    document = Document(path)
    sections: list[ExtractedSection] = []
    current_heading: str | None = None
    current_paragraphs: list[str] = []

    def flush() -> None:
        if current_paragraphs:
            sections.append(
                ExtractedSection(
                    content="\n\n".join(current_paragraphs),
                    heading=current_heading,
                    metadata={"source_type": "docx_section"},
                )
            )

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            flush()
            current_heading = text
            current_paragraphs = [text]
        else:
            current_paragraphs.append(text)

    flush()
    sections.extend(_extract_docx_tables(document))
    return sections


def _extract_docx_tables(document: Document) -> list[ExtractedSection]:
    sections: list[ExtractedSection] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            values: list[str] = []
            for cell in row.cells:
                value = _normalize_cell_text(cell.text)
                if value:
                    values.append(value)
            if values:
                rows.append(" | ".join(values))
        if rows:
            heading = f"Table {table_index}"
            sections.append(
                ExtractedSection(
                    content="\n".join(rows),
                    heading=heading,
                    metadata={
                        "source_type": "docx_table",
                        "table_index": table_index,
                        "row_count": len(rows),
                    },
                )
            )
    return sections


def _normalize_cell_text(text: str) -> str:
    return " ".join(part.strip() for part in text.splitlines() if part.strip())
